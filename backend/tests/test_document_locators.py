"""Locator preservation tests for PDF and DOCX ingestion."""

from __future__ import annotations

import math
import sys
import tempfile
import types
import unittest
from pathlib import Path
from unittest import mock

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from mineru_service.adapter import _content_list_to_blocks
from services.knowledge.document_pipeline import (
    _normalize_ir,
    _parse_docx,
    _parse_pdf_pdfplumber,
)


class MinerULocatorTests(unittest.TestCase):
    def test_content_list_preserves_normalized_page_and_bbox(self):
        blocks = _content_list_to_blocks(
            [
                {
                    "type": "text",
                    "text": "Article 1",
                    "page_idx": 0,
                    "bbox": [100, 200, 300, 400],
                }
            ]
        )

        self.assertIn("locator", blocks[0], "MinerU bbox locator was dropped")
        self.assertEqual(
            blocks[0]["locator"],
            {
                "kind": "pdf",
                "page_number": 1,
                "origin": "top_left",
                "coordinate_system": "normalized_0_1000",
                "rects": [{"x0": 100.0, "y0": 200.0, "x1": 300.0, "y1": 400.0}],
                "precision": "exact",
            },
        )

    def test_content_list_rejects_malformed_exact_bbox(self):
        malformed = ([300, 200, 100, 400], [0, 0, math.nan, 10], [1, 2, 3])

        for bbox in malformed:
            with self.subTest(bbox=bbox):
                block = _content_list_to_blocks(
                    [{"type": "text", "text": "x", "page_idx": 0, "bbox": bbox}]
                )[0]
                self.assertIn("locator", block, "malformed bbox needs page precision")
                self.assertEqual(block["locator"]["precision"], "page")
                self.assertEqual(block["locator"]["rects"], [])


class PdfPlumberLocatorTests(unittest.TestCase):
    def test_fallback_uses_word_geometry_and_keeps_repeated_lines_distinct(self):
        class FakePage:
            width = 200
            height = 400

            def __init__(self):
                self.extract_words_called = False

            def extract_tables(self):
                return []

            def extract_words(self):
                self.extract_words_called = True
                return [
                    {"text": "same", "x0": 20, "x1": 60, "top": 40, "bottom": 60},
                    {"text": "same", "x0": 20, "x1": 60, "top": 80, "bottom": 100},
                ]

            def extract_text(self):
                raise AssertionError("locator fallback must not derive blocks from extract_text")

        page = FakePage()
        fake_pdf = types.SimpleNamespace(pages=[page])
        fake_context = mock.MagicMock()
        fake_context.__enter__.return_value = fake_pdf
        fake_module = types.SimpleNamespace(open=mock.Mock(return_value=fake_context))

        with mock.patch.dict(sys.modules, {"pdfplumber": fake_module}):
            blocks, pages = _parse_pdf_pdfplumber(Path("ignored.pdf"))

        self.assertTrue(page.extract_words_called)
        self.assertEqual(pages, 1)
        self.assertEqual([block["text"] for block in blocks], ["same", "same"])
        self.assertNotEqual(blocks[0]["locator"]["rects"], blocks[1]["locator"]["rects"])
        self.assertEqual(
            blocks[0]["locator"],
            {
                "kind": "pdf",
                "page_number": 1,
                "origin": "top_left",
                "coordinate_system": "normalized_0_1000",
                "rects": [{"x0": 100.0, "y0": 100.0, "x1": 300.0, "y1": 150.0}],
                "precision": "exact",
            },
        )

    def test_ir_keeps_locator_and_existing_callers_fields(self):
        locator = {
            "kind": "pdf",
            "page_number": 2,
            "origin": "top_left",
            "coordinate_system": "normalized_0_1000",
            "rects": [{"x0": 1.0, "y0": 2.0, "x1": 3.0, "y1": 4.0}],
            "precision": "exact",
        }
        ir = _normalize_ir(
            doc_id="doc",
            title="Title",
            filename="doc.pdf",
            mime="application/pdf",
            pages=2,
            raw_blocks=[{"type": "paragraph", "text": "text", "page": 2, "locator": locator}],
        )

        block = ir["blocks"][0]
        self.assertIn("locator", block, "normalization dropped the raw locator")
        self.assertEqual(block["locator"], locator)
        self.assertEqual(block["page_start"], 2)
        self.assertEqual(block["page_end"], 2)
        self.assertEqual(block["block_id"], "b1")


class DocxLocatorTests(unittest.TestCase):
    def test_xml_body_order_cell_identity_and_unicode_ranges_are_stable(self):
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - dependency is required by fallback
            self.skipTest(str(exc))

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "ordered.docx"
            doc = Document()
            doc.add_paragraph("A\U0001f600B")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "duplicate"
            table.cell(0, 1).text = "duplicate"
            doc.add_paragraph("duplicate")
            doc.save(path)

            first, pages, engine = _parse_docx(path)
            second, _, _ = _parse_docx(path)

        self.assertIsNone(pages)
        self.assertEqual(engine, "python-docx")
        self.assertEqual(first, second)
        self.assertEqual([block["type"] for block in first], ["paragraph", "table", "paragraph"])
        self.assertNotIn("page", first[0])
        self.assertEqual(first[0]["locator"]["container_kind"], "paragraph")
        self.assertEqual(
            first[0]["locator"]["text_range"],
            {"start": 0, "end": 3, "unit": "unicode_code_point"},
        )

        cell_locators = first[1]["locators"]
        duplicate_locators = cell_locators + [first[2]["locator"]]
        self.assertEqual(
            [locator["container_kind"] for locator in duplicate_locators],
            ["table_cell", "table_cell", "paragraph"],
        )
        self.assertEqual(
            [locator["document_order"] for locator in duplicate_locators],
            [1, 2, 3],
        )
        self.assertEqual(len({locator["locator_id"] for locator in duplicate_locators}), 3)
        self.assertTrue(all(locator["block_id"] for locator in duplicate_locators))


if __name__ == "__main__":
    unittest.main()
